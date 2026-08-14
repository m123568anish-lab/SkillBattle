"""
=========================================================

SkillBattle

DOCX Resume Parser

Uses python-docx to extract text.

=========================================================
"""

from __future__ import annotations

from pathlib import Path

from docx import Document


class DOCXParser:

    """
    DOCX Resume Parser
    """

    def extract_text(
        self,
        docx_path: str,
    ) -> str:

        path = Path(docx_path)

        if not path.exists():

            raise FileNotFoundError(
                f"File not found: {docx_path}"
            )

        document = Document(docx_path)

        paragraphs: list[str] = []

        for paragraph in document.paragraphs:

            text = paragraph.text.strip()

            if text:

                paragraphs.append(text)

        return "\n".join(paragraphs)

    # --------------------------------------------------

    def paragraph_count(
        self,
        docx_path: str,
    ) -> int:

        document = Document(docx_path)

        return len(document.paragraphs)

    # --------------------------------------------------

    def table_count(
        self,
        docx_path: str,
    ) -> int:

        document = Document(docx_path)

        return len(document.tables)

    # --------------------------------------------------

    def metadata(
        self,
        docx_path: str,
    ) -> dict:

        document = Document(docx_path)

        props = document.core_properties

        return {

            "author": props.author,

            "title": props.title,

            "subject": props.subject,

            "keywords": props.keywords,

            "created": props.created,

            "modified": props.modified,

        }


docx_parser = DOCXParser()